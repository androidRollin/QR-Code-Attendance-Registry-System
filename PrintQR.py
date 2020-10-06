import docx
import csv
from docx2pdf import convert
import os

def QRstoPDF(path):
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run()
    try:
        with open(path, "r") as csv_file:
            csv_reader = csv.reader(csv_file)
            next(csv_reader)
            count = 0
            for line in csv_reader:
                count+=1
                FileNameQR = "QRimages/"+line[0]+"QR.jpg"
                Section = line[2]
                r.add_picture(FileNameQR, width=docx.shared.Inches(3), height=docx.shared.Inches(3))
                if (count == 6):
                    p = doc.add_paragraph()
                    r = p.add_run()
                    count = 0
    except FileNotFoundError:
        print("FileNOtFound")

    pdfname = "C:/Users/ROLINE JOHN AGUILAR/Downloads/"+line[2]+"QRs"+".pdf"
    wordname = "QRGenerated.docx"
    doc.save(wordname)

    convert(wordname, pdfname)
    os.remove(wordname)
    os.startfile(pdfname)
    os.remove("QRData.csv")




