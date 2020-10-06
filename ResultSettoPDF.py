import docx
from docx2pdf import convert
import os
from ProjectAnalytics import getfeatures_test,getfeatures_train,gettarget_test,gettarget_train
from datetime import datetime

def TabletoPDF(text):
    date = str(datetime.now().strftime('%d-%m-%y'))
    doc = docx.Document()
    p = doc.add_paragraph()
    runner = p.add_run("Data Analytics " + date)
    runner.bold = True
    p = doc.add_paragraph(text)

    pdfname = "C:/Users/ROLINE JOHN AGUILAR/Downloads/DataAnalyticsTables.pdf"
    wordname = "DA.docx"
    doc.save(wordname)

    convert(wordname, pdfname)
    os.remove(wordname)
    os.startfile(pdfname)

def getFeaturesattr(features_train,features_test,target_train,target_test):
    LongResult = " Training Features \n" + str(features_train) + "\n\n Testing Features \n" + str(features_test) +  \
                 "\n\n Training Target \n" + str(target_train) + "\n\n Testing Target \n" + str(target_test) +"\n"
    return str(LongResult)

text = getFeaturesattr(getfeatures_train(),getfeatures_test(),gettarget_train(),gettarget_test())
TabletoPDF(text)